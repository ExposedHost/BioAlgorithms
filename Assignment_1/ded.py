from docx import Document
from docx.shared import RGBColor
from Bio import SeqIO

def highlight_diff(string1, string2,i):
    paragraph = doc.add_paragraph()
    if(i==1):
        run = paragraph.add_run(string1)
    doc.add_paragraph()  # Add an empty paragraph for spacing

    for char1, char2 in zip(string1, string2):
        run = paragraph.add_run(char2)
        if char1 != char2:
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)  # Red color

    # If the second string is longer than the first, append the remaining characters
    if len(string2) > len(string1):
        remaining_text = string2[len(string1):]
        paragraph.add_run(remaining_text)

def read_fasta(filename):
    for record in SeqIO.parse(filename, "fasta"):
        return str(record.seq)

# Example usage:
doc = Document()
filename1 = "seq2.fasta"
string1 = read_fasta(filename1)

for i in range(1, 10):
    doc.add_heading(f"Mutant {i}", level=1)
    filename2 = f"mutant_seq2_{i}.fasta"
    string2 = read_fasta(filename2)
    highlight_diff(string1, string2,i)
    # doc.add_page_break()

doc.save("highlighted_text2.docx")
